#!/usr/bin/env python3
"""Generate cabinet immobilier letter from template.

Usage:
    generate.py <data.json> <output.docx>      -> DOCX
    generate.py <data.json> <output.pdf>       -> DOCX intermediaire puis PDF
                                                  (conversion LibreOffice, mise
                                                   en page et accents preserves)

La generation PDF NE construit JAMAIS le PDF a la main : elle part du DOCX
produit par docxtpl (template rempli) et le convertit via soffice. C'est la
seule methode fiable.
"""
import sys
import os
import json
import shutil
import subprocess
from docxtpl import DocxTemplate
from docx import Document

TEMPLATE_PATH = os.environ.get(
    "TEMPLATE_COURRIER",
    os.path.join(os.path.dirname(os.path.abspath(__file__)), "template.docx"),
)


def create_template(template_path):
    """Create base template if not exists"""
    doc = Document()
    doc.add_heading("{{nom_cabinet}}", 0)
    doc.add_paragraph("{{adresse_cabinet}}")
    doc.add_paragraph()
    doc.add_paragraph("{{nom_destinataire}}")
    doc.add_paragraph("{{adresse_destinataire}}")
    doc.add_paragraph()
    doc.add_paragraph("Le {{date}}")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Objet : ").bold = True
    p.add_run("{{objet}}")
    doc.add_paragraph()
    doc.add_paragraph("Madame, Monsieur,")
    doc.add_paragraph()
    doc.add_paragraph("{{corps}}")
    doc.add_paragraph()
    doc.add_paragraph("Veuillez agreer, Madame, Monsieur, l\u0027expression de mes salutations distinguees.")
    doc.add_paragraph()
    doc.add_paragraph("{{signataire}}")
    doc.add_paragraph("{{fonction_signataire}}")
    doc.save(template_path)


from xml.sax.saxutils import escape as _xml_escape

def _assainir(data):
    """Le corps metier peut porter <, > ou & : docxtpl l'injecte tel quel dans le XML
    du template -> .docx CORROMPU (constate 18/07). On echappe le XML pour toutes les
    valeurs et on convertit les sauts de ligne du corps en sauts de ligne XML souples,
    compatibles avec un paragraphe simple du template ({{corps}}).
    NB : l'echappement suppose autoescape desactive (defaut docxtpl).
    """
    out = {}
    for k, v in (data or {}).items():
        s = "" if v is None else str(v)
        s = _xml_escape(s)
        s = s.replace("\n", '<w:br/>')
        out[k] = s
    return out

GABARITS_PATH = os.environ.get("GABARITS_PATH",
                               os.path.join(os.path.dirname(os.path.abspath(__file__)),
                                            "..", "..", "..", "..", "data-seed",
                                            "gabarits_motif.json"))


def render_md(data, md_path):
    """Markdown : ni LibreOffice ni docxtpl. La mise en page vient du fichier de
    gabarits (gabarit_md), comme celle du Word vient de template.docx -- la forme
    est une donnee, quel que soit le format."""
    modele = None
    try:
        with open(GABARITS_PATH, encoding="utf-8") as fh:
            modele = json.load(fh).get("gabarit_md")
    except Exception:
        modele = None
    if not modele:
        modele = ("**{nom_cabinet}**\n{adresse_cabinet}\n\n{nom_destinataire}\n"
                  "{adresse_destinataire}\n\nLe {date}\n\n**Objet : {objet}**\n\n"
                  "Madame, Monsieur,\n\n{corps}\n\n{signataire}\n"
                  "{fonction_signataire}\n")

    class _Vide(dict):
        def __missing__(self, k):
            return ""

    with open(md_path, "w", encoding="utf-8") as fh:
        fh.write(modele.format_map(_Vide(data or {})))


def render_docx(data, docx_path):
    """Rend le template docxtpl vers docx_path."""
    if not os.path.exists(TEMPLATE_PATH):
        create_template(TEMPLATE_PATH)
    template = DocxTemplate(TEMPLATE_PATH)
    template.render(_assainir(data))
    template.save(docx_path)


def convert_to_pdf(docx_path, pdf_path):
    """Convertit un docx en pdf via LibreOffice. Retourne le chemin du PDF.

    Part du docx (deja bien forme) -> mise en page + accents preserves.
    """
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise RuntimeError(
            "LibreOffice (soffice) absent du container : conversion PDF impossible. "
            "Installer libreoffice-writer-nogui dans l'image du relay."
        )
    outdir = os.path.dirname(os.path.abspath(pdf_path)) or "."
    # LibreOffice impose le nom <base>.pdf dans outdir ; on aligne ensuite si besoin
    subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", outdir, docx_path],
        check=True, capture_output=True, timeout=120,
    )
    produced = os.path.join(outdir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
    if produced != os.path.abspath(pdf_path) and os.path.exists(produced):
        shutil.move(produced, pdf_path)
    if not os.path.exists(pdf_path) or os.path.getsize(pdf_path) < 1024:
        raise RuntimeError("Conversion PDF echouee ou PDF anormalement petit.")
    return pdf_path


def main():
    if len(sys.argv) != 3:
        print("Usage: generate.py <data.json> <output.docx|output.pdf>")
        sys.exit(1)

    data_file = sys.argv[1]
    output_file = sys.argv[2]
    want_pdf = output_file.lower().endswith(".pdf")
    want_md  = output_file.lower().endswith(".md")

    with open(data_file, encoding="utf-8") as f:
        data = json.load(f)

    if want_md:
        render_md(data, output_file)
        print(f"OK - {os.path.getsize(output_file)} bytes (Markdown)")
    elif want_pdf:
        # docx intermediaire a cote du pdf demande
        docx_path = os.path.splitext(output_file)[0] + ".docx"
        render_docx(data, docx_path)
        convert_to_pdf(docx_path, output_file)
        print(f"OK - {os.path.getsize(output_file)} bytes (PDF via LibreOffice, docx: {docx_path})")
    else:
        render_docx(data, output_file)
        print(f"OK - {os.path.getsize(output_file)} bytes")


if __name__ == "__main__":
    main()
